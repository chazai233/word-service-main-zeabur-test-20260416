"""
Word Document Generation Service - Precision Formatting Version (v8.0)
修复重点：
1. 消除大标题(1、)的错误缩进。
2. 实现"人员投入"等关键词的【局部加粗】（而非整行）。
3. 严格控制子标题((1))和统计项的首行缩进。
"""

import base64
import io
import json
import re
import os
import requests
from collections import Counter
from copy import deepcopy
from datetime import datetime, timedelta, date
from zoneinfo import ZoneInfo
from urllib.parse import urlparse, parse_qs
from typing import List, Optional, Dict, Any
from enum import Enum

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI(title="Word Service Precision", version="8.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CN_TEMPLATE_PATH = os.path.join(BASE_DIR, "templates", "cn-template.docx")
EN_TEMPLATE_PATH = os.path.join(BASE_DIR, "templates", "en-template.docx")
BANGKOK_TZ = ZoneInfo("Asia/Bangkok")
PAKBENG_COORD_FALLBACK = (19.8933, 101.1348)

# ---------------- 核心排版逻辑 ----------------

def format_run_font(run, size=10.5, bold=False):
    """统一设置字体格式"""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold

def process_and_add_line(cell, line_text):
    """
    智能处理每一行的格式：缩进、加粗、分割
    """
    line_text = line_text.strip()
    if not line_text: return

    # 创建新段落（注意：不使用 add_run("\n") 而是 add_paragraph 以便单独控制每一行的缩进）
    # 如果是单元格刚清空后的第一个段落，直接使用，否则新建
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    # --- 1. 规则匹配 ---
    
    # 规则A：大标题 (例如 "1、右岸施工营地")
    # 特征：数字开头 + 顿号或点
    if re.match(r"^\d+[、\.]", line_text):
        p.paragraph_format.first_line_indent = Pt(0) # 【关键】强制不缩进
        run = p.add_run(line_text)
        format_run_font(run, bold=True) # 大标题整行加粗
        return

    # 规则B：统计项 (例如 "人员投入：...")
    # 特征：包含特定关键词
    keywords = ["人员投入", "设备投入", "累计工程量"]
    hit_keyword = None
    for kw in keywords:
        if kw in line_text:
            hit_keyword = kw
            break
    
    if hit_keyword:
        p.paragraph_format.first_line_indent = Pt(24) # 【关键】强制缩进 2 字符
        
        # 【局部加粗逻辑】
        # 将文本切分为两部分：关键词前缀(加粗) + 剩余内容(不加粗)
        # 例如 "人员投入：张三" -> "人员投入：" (粗) + " 张三" (细)
        
        # 尝试找到冒号的位置
        split_index = -1
        if "：" in line_text:
            split_index = line_text.index("：") + 1
        elif ":" in line_text:
            split_index = line_text.index(":") + 1
        else:
            # 如果没有冒号，就只加粗关键词本身
            split_index = line_text.index(hit_keyword) + len(hit_keyword)
            
        prefix = line_text[:split_index]
        content = line_text[split_index:]
        
        # 写入前缀（加粗）
        run1 = p.add_run(prefix)
        format_run_font(run1, bold=True)
        
        # 写入内容（正常）
        run2 = p.add_run(content)
        format_run_font(run2, bold=False)
        return

    # 规则C：子标题 / 具体内容 (例如 "(1) 场地精平")
    # 特征：以 (数字) 或 （数字） 开头
    if re.match(r"^[\(（]\d+[\)）]", line_text):
        p.paragraph_format.first_line_indent = Pt(24) # 【关键】强制缩进 2 字符
        run = p.add_run(line_text)
        format_run_font(run, bold=False) # 内容不加粗
        return

    # 规则D：其他普通文本
    # 默认缩进2字符（因为通常是正文延续），或者0？
    # 根据你的截图，如果不符合上述规则，通常是正文描述，建议缩进2字符对齐
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(line_text)
    format_run_font(run, bold=False)

# ---------------- 辅助函数 ----------------

def find_target_table(doc: Document, index: int) -> Optional[Any]:
    if 0 <= index < len(doc.tables):
        return doc.tables[index]
    return None

def update_table_row(table, row_name: str, today: str, total: str):
    """表格行更新逻辑"""
    if not table.rows: return
    
    name_col = 1
    cols_count = len(table.rows[0].cells)
    today_col = 4 if cols_count > 4 else cols_count - 2
    total_col = 5 if cols_count > 5 else cols_count - 1
    
    for row in table.rows:
        if len(row.cells) <= max(name_col, today_col, total_col): continue
        cell_text = row.cells[name_col].text.strip()
        if row_name in cell_text: 
            # 填入数字时也应用字体规范
            if today and today != "-":
                cell = row.cells[today_col]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(today))
                format_run_font(run)
            if total and total != "-":
                cell = row.cells[total_col]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(total))
                format_run_font(run)
            return


DAILY_STATS_HEADER_CANDIDATES: Dict[str, List[str]] = {
    "seq": ["序号", "S/N", "SN"],
    "location": ["施工部位", "Construction Area"],
    "content": ["施工内容", "Activities", "Activity"],
    "quantity": ["日完成量", "完成工程量", "Quantities Completed", "Planned Qty.", "Qty."],
    "remarks": ["备注", "Remarks"],
}


def normalize_header_text(text: str) -> str:
    # 统一去空白并小写，便于中英表头做 contains 匹配
    return re.sub(r"\s+", "", text or "").lower()


def find_daily_stats_table(doc: Document):
    """
    自动识别“当日施工统计表”：
    需在同一表头行识别到：
    - 中文：序号、施工部位、施工内容、日完成量、备注
    - 英文：S/N、Construction Area、Activities、Quantities Completed、Remarks
    """
    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            header_map: Dict[str, int] = {}
            for col_index, cell in enumerate(row.cells):
                cell_text = normalize_header_text(cell.text)
                for canonical_key, candidates in DAILY_STATS_HEADER_CANDIDATES.items():
                    if canonical_key in header_map:
                        continue
                    for candidate in candidates:
                        if normalize_header_text(candidate) in cell_text:
                            header_map[canonical_key] = col_index
                            break
            if len(header_map) == len(DAILY_STATS_HEADER_CANDIDATES):
                return table, row_index, header_map
    return None, None, None


def pick_first_value(item: Dict[str, Any], keys: List[str], default: str = "") -> str:
    for key in keys:
        value = item.get(key)
        if value is not None:
            text = str(value).strip()
            if text:
                return text
    return default


def normalize_quantity_text(quantity: str, content: str = "") -> str:
    q = (quantity or "").strip()
    if not q:
        return q

    # 先处理常见写法
    q = q.replace("㎡", "m²").replace("m2", "m²").replace("m^2", "m²")
    q = q.replace("m3", "m³").replace("m^3", "m³")

    return q


def detect_suspect_quantity(items: List[Dict[str, str]]) -> List[int]:
    """
    严格模式：禁止对损坏单位做猜测修复。
    若出现 m? / m� 等异常单位，返回对应行号（1-based）。
    """
    bad_rows: List[int] = []
    for i, item in enumerate(items, start=1):
        q = (item.get("quantity") or "").strip()
        if re.search(r"m[\?\uFFFD]", q):
            bad_rows.append(i)
    return bad_rows


def normalize_daily_stats_items(raw_items: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    def normalize_text(text: str) -> str:
        text = "" if text is None else str(text)
        # 兼容上游把中文以 "\\u4e2d\\u6587" 字符串形式传入
        if "\\u" in text and not re.search(r"[\u4e00-\u9fff]", text):
            try:
                text = bytes(text, "utf-8").decode("unicode_escape")
            except Exception:
                pass
        return text.strip()

    normalized: List[Dict[str, str]] = []
    for raw in raw_items:
        if not isinstance(raw, dict):
            continue
        seq_text = normalize_text(pick_first_value(raw, ["seq", "序号", "sn", "serial"], ""))
        location_text = normalize_text(pick_first_value(raw, ["location", "施工部位", "area"], ""))
        content_text = normalize_text(pick_first_value(raw, ["content", "施工内容", "activity"], ""))
        quantity_text = normalize_text(pick_first_value(raw, ["quantity", "日完成量", "完成工程量", "qty"], ""))
        remarks_text = normalize_text(pick_first_value(raw, ["remarks", "备注", "remark"], ""))

        row = {
            "seq": seq_text,
            "location": location_text,
            "content": content_text,
            "quantity": normalize_quantity_text(quantity_text, content_text),
            "remarks": remarks_text,
        }
        # 施工内容为空的数据直接忽略
        if row["content"]:
            normalized.append(row)
    return normalized


def parse_daily_stats_from_base64(encoded: str) -> Optional[List[Dict[str, str]]]:
    if not encoded:
        return None
    try:
        raw = base64.b64decode(encoded)
        text = raw.decode("utf-8")
        data = json.loads(text)
    except Exception:
        return None
    if not isinstance(data, list):
        return None
    return normalize_daily_stats_items(data)


def detect_garbled_daily_stats(items: List[Dict[str, str]]) -> bool:
    """
    识别常见乱码损坏场景：
    - 关键字段出现连续 ???
    - 且全文没有任何中文字符（通常代表中文已在上游丢失为 ?）
    """
    if not items:
        return False

    has_cjk = False
    has_suspect = False
    for item in items:
        text = " ".join([item.get("location", ""), item.get("content", ""), item.get("quantity", "")])
        if re.search(r"[\u4e00-\u9fff]", text):
            has_cjk = True
        if "�" in text or re.search(r"\?{2,}", text):
            has_suspect = True
    return has_suspect and not has_cjk


def parse_daily_stats_from_content(content: str) -> Optional[List[Dict[str, str]]]:
    """
    兼容：当 content 直接传入 JSON 数组字符串时，自动识别并切换到统计表模式。
    """
    if not isinstance(content, str):
        return None
    text = content.strip()
    if not text.startswith("["):
        return None
    try:
        data = json.loads(text)
    except Exception:
        return None
    if not isinstance(data, list):
        return None
    return normalize_daily_stats_items(data)


def parse_daily_stats_json_text(text: Optional[str]) -> Optional[List[Dict[str, str]]]:
    if not isinstance(text, str):
        return None
    raw = text.strip()
    if not raw:
        return None
    try:
        data = json.loads(raw)
    except Exception:
        return None
    if not isinstance(data, list):
        return None
    return normalize_daily_stats_items(data)


def parse_english_translated_items(raw_text: Optional[str]) -> Optional[List[Dict[str, str]]]:
    if not isinstance(raw_text, str):
        return None
    raw = raw_text.strip()
    if not raw:
        return None

    data_obj: Any = None
    try:
        data_obj = json.loads(raw)
    except Exception:
        match = re.search(r"\{[\s\S]*\}", raw)
        if match:
            try:
                data_obj = json.loads(match.group())
            except Exception:
                return None
        else:
            return None

    translated = None
    if isinstance(data_obj, dict):
        translated = data_obj.get("translated_data")
    elif isinstance(data_obj, list):
        translated = data_obj

    if not isinstance(translated, list):
        return None

    items: List[Dict[str, str]] = []
    for row in translated:
        if not isinstance(row, dict):
            continue
        seq_text = str(row.get("seq", "")).strip()
        location_text = str(row.get("location_en") or row.get("location") or "").strip()
        content_text = str(row.get("content_en") or row.get("content") or "").strip()
        quantity_raw = str(row.get("quantity_en") or row.get("quantity") or "").strip()
        remarks_text = str(
            row.get("remarks_en") or row.get("remarks") or row.get("shift") or ""
        ).strip()

        if not content_text:
            continue

        items.append(
            {
                "seq": seq_text,
                "location": location_text,
                "content": content_text,
                "quantity": normalize_quantity_text(quantity_raw, content_text),
                "remarks": remarks_text,
            }
        )
    return items or None


def _resolve_business_date_by_trigger_time(trigger_dt: datetime) -> date:
    if trigger_dt.tzinfo is None:
        local_dt = trigger_dt.replace(tzinfo=BANGKOK_TZ)
    else:
        local_dt = trigger_dt.astimezone(BANGKOK_TZ)
    if local_dt.hour >= 18:
        return local_dt.date()
    return (local_dt - timedelta(days=1)).date()


def parse_trigger_date(value: Optional[str]) -> date:
    if isinstance(value, str):
        text = value.strip()
        if text:
            # 日期字符串视为手工指定业务日期，保持兼容
            for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
                try:
                    parsed_date = datetime.strptime(text, fmt).date()
                    now_local = datetime.now(BANGKOK_TZ)
                    # 兼容 Dify 仅传 YYYY-MM-DD（通常为“今天”）的场景：
                    # 若与服务端“今天”一致，则按触发时刻走 18:00 分界规则。
                    if parsed_date == now_local.date():
                        return _resolve_business_date_by_trigger_time(now_local)
                    return parsed_date
                except Exception:
                    pass

            normalized = text
            if normalized.endswith("Z"):
                normalized = normalized[:-1] + "+00:00"
            try:
                parsed_dt = datetime.fromisoformat(normalized)
                return _resolve_business_date_by_trigger_time(parsed_dt)
            except Exception:
                pass

            for fmt in (
                "%Y-%m-%d %H:%M:%S",
                "%Y-%m-%d %H:%M",
                "%Y/%m/%d %H:%M:%S",
                "%Y/%m/%d %H:%M",
                "%Y.%m.%d %H:%M:%S",
                "%Y.%m.%d %H:%M",
            ):
                try:
                    parsed_dt = datetime.strptime(text, fmt).replace(tzinfo=BANGKOK_TZ)
                    return _resolve_business_date_by_trigger_time(parsed_dt)
                except Exception:
                    pass

    return _resolve_business_date_by_trigger_time(datetime.now(BANGKOK_TZ))


def format_cn_date(d: date) -> str:
    return f"{d.year}年{d.month}月{d.day}日"


def format_en_date(d: date) -> str:
    months = ["Jan.", "Feb.", "Mar.", "Apr.", "May", "Jun.", "Jul.", "Aug.", "Sep.", "Oct.", "Nov.", "Dec."]
    return f"{months[d.month - 1]} {d.day}, {d.year}"


_EN_MONTH_REGEX = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\.?\s*\d{1,2},\s*\d{4}"
_MONTH_DOT_DAY_REGEX = r"(?<!\d)(?:0?[1-9]|1[0-2])\s*\.\s*(?:0?[1-9]|[12]\d|3[01])(?!\d)"
_SIGNATURE_DATE_LABEL_REGEX = re.compile(r"(?:^|\s)(?:日期|Date)\s*[:：]", flags=re.IGNORECASE)


def format_slash_date(d: date) -> str:
    return f"{d.year}/{d.month}/{d.day}"


def format_month_dot_day(d: date) -> str:
    return f"{d.month}.{d.day}"


def replace_date_in_text(text: str, target_day: date, allow_month_dot_day: bool = False) -> tuple[str, bool]:
    source = "" if text is None else str(text)
    updated = source

    updated = re.sub(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日", format_cn_date(target_day), updated)
    updated = re.sub(r"\d{4}\s*/\s*\d{1,2}\s*/\s*\d{1,2}", format_slash_date(target_day), updated)
    updated = re.sub(_EN_MONTH_REGEX, format_en_date(target_day), updated, flags=re.IGNORECASE)
    updated = re.sub(r"\d{4}-\d{1,2}-\d{1,2}", target_day.isoformat(), updated)
    if allow_month_dot_day:
        updated = re.sub(_MONTH_DOT_DAY_REGEX, format_month_dot_day(target_day), updated)

    return updated, updated != source


def _get_pakbeng_coordinates_openweather(api_key: str) -> tuple[float, float]:
    try:
        resp = requests.get(
            "https://api.openweathermap.org/geo/1.0/direct",
            params={"q": "Pak Beng,LA", "limit": 1, "appid": api_key},
            timeout=20,
        )
        resp.raise_for_status()
        data = resp.json()
        if isinstance(data, list) and data:
            lat = data[0].get("lat")
            lon = data[0].get("lon")
            if lat is not None and lon is not None:
                return float(lat), float(lon)
    except Exception:
        pass
    return PAKBENG_COORD_FALLBACK


def _openweather_condition_text(weather_id: Optional[int]) -> tuple[str, str]:
    if weather_id is None:
        return "未知", "Unknown"
    exact_map = {
        800: ("晴", "Clear"),
        801: ("少云", "Few Clouds"),
        802: ("多云", "Scattered Clouds"),
        803: ("阴", "Broken Clouds"),
        804: ("阴天", "Overcast Clouds"),
    }
    if weather_id in exact_map:
        return exact_map[weather_id]
    category_map = {
        2: ("雷暴", "Thunderstorm"),
        3: ("毛毛雨", "Drizzle"),
        5: ("雨", "Rain"),
        6: ("雪", "Snow"),
        7: ("雾", "Mist"),
    }
    return category_map.get(weather_id // 100, ("未知", "Unknown"))


def _build_weather_entry(
    *,
    dt_ts: int,
    weather_id: Optional[int],
    temp_min: Optional[float],
    temp_max: Optional[float],
    wind_speed_ms: Optional[float],
) -> Dict[str, Any]:
    return {
        "local_dt": datetime.fromtimestamp(dt_ts, tz=BANGKOK_TZ),
        "weather_id": weather_id,
        "temp_min": temp_min,
        "temp_max": temp_max,
        "wind_kmh": (wind_speed_ms * 3.6) if wind_speed_ms is not None else None,
    }


def _to_int(value: Any) -> Optional[int]:
    try:
        return int(value)
    except Exception:
        return None


def _to_float(value: Any) -> Optional[float]:
    try:
        return float(value)
    except Exception:
        return None


def _summarize_weather_entries(entries: List[Dict[str, Any]], target_date: date) -> Dict[str, str]:
    same_day = [x for x in entries if x.get("local_dt") and x["local_dt"].date() == target_date]
    rows = same_day if same_day else entries

    weather_id = None
    noon_rows = [x for x in rows if x.get("local_dt") and x["local_dt"].hour == 12 and isinstance(x.get("weather_id"), int)]
    if noon_rows:
        weather_id = noon_rows[0]["weather_id"]
    else:
        ids = [x["weather_id"] for x in rows if isinstance(x.get("weather_id"), int)]
        if ids:
            weather_id = Counter(ids).most_common(1)[0][0]

    temp_min_values = [x["temp_min"] for x in rows if isinstance(x.get("temp_min"), (int, float))]
    temp_max_values = [x["temp_max"] for x in rows if isinstance(x.get("temp_max"), (int, float))]
    wind_values = [x["wind_kmh"] for x in rows if isinstance(x.get("wind_kmh"), (int, float))]

    weather_zh, weather_en = _openweather_condition_text(weather_id)
    wind_zh, wind_en = _wind_force_text(max(wind_values) if wind_values else None)
    return {
        "weather_zh": weather_zh,
        "weather_en": weather_en,
        "wind_zh": wind_zh,
        "wind_en": wind_en,
        "temp": _temp_text(min(temp_min_values) if temp_min_values else None, max(temp_max_values) if temp_max_values else None),
    }


def _wind_force_text(speed_kmh: Optional[float]) -> tuple[str, str]:
    if speed_kmh is None:
        return "未知", "Unknown"
    steps = [
        (1, "无风", "Calm"),
        (6, "软风", "Light Air"),
        (12, "轻风", "Light Breeze"),
        (20, "微风", "Gentle Breeze"),
        (29, "和风", "Moderate Breeze"),
        (39, "清劲风", "Fresh Breeze"),
        (50, "强风", "Strong Breeze"),
        (62, "疾风", "High Wind"),
        (75, "大风", "Gale"),
        (89, "烈风", "Strong Gale"),
        (103, "狂风", "Storm"),
        (118, "暴风", "Violent Storm"),
    ]
    for bound, zh, en in steps:
        if speed_kmh < bound:
            return zh, en
    return "飓风", "Hurricane"


def _temp_text(min_temp: Optional[float], max_temp: Optional[float]) -> str:
    if min_temp is None or max_temp is None:
        return "--"
    return f"{round(min_temp)}℃/{round(max_temp)}℃"


def fetch_pakbeng_weather(target_date: date) -> tuple[Dict[str, str], Optional[str]]:
    weather = {
        "weather_zh": "未知",
        "weather_en": "Unknown",
        "wind_zh": "未知",
        "wind_en": "Unknown",
        "temp": "--",
    }
    api_key = os.getenv("OPENWEATHER_API_KEY", "").strip()
    if not api_key:
        return weather, "未配置 OPENWEATHER_API_KEY，已使用默认天气"

    try:
        lat, lon = _get_pakbeng_coordinates_openweather(api_key)
        today = datetime.now(BANGKOK_TZ).date()
        entries: List[Dict[str, Any]] = []
        warnings: List[str] = []

        if target_date < today:
            history_ts = int(datetime(target_date.year, target_date.month, target_date.day, 12, 0, tzinfo=BANGKOK_TZ).timestamp())
            try:
                resp = requests.get(
                    "https://api.openweathermap.org/data/3.0/onecall/timemachine",
                    params={
                        "lat": lat,
                        "lon": lon,
                        "dt": history_ts,
                        "appid": api_key,
                        "units": "metric",
                    },
                    timeout=25,
                )
                resp.raise_for_status()
                payload = resp.json()
                for item in payload.get("data") or []:
                    dt_ts = _to_int(item.get("dt"))
                    if dt_ts is None:
                        continue
                    if datetime.fromtimestamp(dt_ts, tz=BANGKOK_TZ).date() != target_date:
                        continue
                    weather_info = (item.get("weather") or [{}])[0]
                    entries.append(
                        _build_weather_entry(
                            dt_ts=dt_ts,
                            weather_id=_to_int(weather_info.get("id")),
                            temp_min=_to_float(item.get("temp")),
                            temp_max=_to_float(item.get("temp")),
                            wind_speed_ms=_to_float(item.get("wind_speed")),
                        )
                    )
            except Exception as exc:
                warnings.append(f"OpenWeather 历史天气查询失败: {exc}")

        if not entries and target_date >= today:
            try:
                resp = requests.get(
                    "https://api.openweathermap.org/data/2.5/forecast",
                    params={
                        "lat": lat,
                        "lon": lon,
                        "appid": api_key,
                        "units": "metric",
                    },
                    timeout=25,
                )
                resp.raise_for_status()
                payload = resp.json()
                for item in payload.get("list") or []:
                    dt_ts = _to_int(item.get("dt"))
                    if dt_ts is None:
                        continue
                    if datetime.fromtimestamp(dt_ts, tz=BANGKOK_TZ).date() != target_date:
                        continue
                    weather_info = (item.get("weather") or [{}])[0]
                    main_info = item.get("main") or {}
                    wind_info = item.get("wind") or {}
                    entries.append(
                        _build_weather_entry(
                            dt_ts=dt_ts,
                            weather_id=_to_int(weather_info.get("id")),
                            temp_min=_to_float(main_info.get("temp_min")),
                            temp_max=_to_float(main_info.get("temp_max")),
                            wind_speed_ms=_to_float(wind_info.get("speed")),
                        )
                    )
            except Exception as exc:
                warnings.append(f"OpenWeather 预报天气查询失败: {exc}")

        if not entries:
            resp = requests.get(
                "https://api.openweathermap.org/data/2.5/weather",
                params={
                    "lat": lat,
                    "lon": lon,
                    "appid": api_key,
                    "units": "metric",
                },
                timeout=20,
            )
            resp.raise_for_status()
            payload = resp.json()
            dt_ts = _to_int(payload.get("dt")) or int(datetime.now(BANGKOK_TZ).timestamp())
            weather_info = (payload.get("weather") or [{}])[0]
            main_info = payload.get("main") or {}
            wind_info = payload.get("wind") or {}
            entries.append(
                _build_weather_entry(
                    dt_ts=dt_ts,
                    weather_id=_to_int(weather_info.get("id")),
                    temp_min=_to_float(main_info.get("temp_min")),
                    temp_max=_to_float(main_info.get("temp_max")),
                    wind_speed_ms=_to_float(wind_info.get("speed")),
                )
            )
            if datetime.fromtimestamp(dt_ts, tz=BANGKOK_TZ).date() != target_date:
                warnings.append("目标日期无可用天气数据，已使用最新实况天气")

        weather = _summarize_weather_entries(entries, target_date)
        warning_text = "；".join(dict.fromkeys(warnings)) if warnings else None
        return weather, warning_text
    except Exception as e:
        return weather, f"天气获取失败，已使用默认值: {e}"

def _parse_bitable_date(value: Any) -> Optional[date]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        ts = float(value)
        if ts > 1_000_000_000_000:
            ts = ts / 1000.0
        try:
            return datetime.fromtimestamp(ts, tz=BANGKOK_TZ).date()
        except Exception:
            return None
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return None
        normalized = raw.replace("年", "-").replace("月", "-").replace("日", "").replace("/", "-").replace(".", "-")
        for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S"):
            try:
                return datetime.strptime(normalized, fmt).date()
            except Exception:
                pass
        try:
            return datetime.fromisoformat(raw.replace("Z", "+00:00")).astimezone(BANGKOK_TZ).date()
        except Exception:
            return None
    if isinstance(value, list):
        for item in value:
            d = _parse_bitable_date(item)
            if d:
                return d
        return None
    if isinstance(value, dict):
        for key in ("value", "text", "name", "date"):
            if key in value:
                d = _parse_bitable_date(value.get(key))
                if d:
                    return d
    return None


def _parse_water_level_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (int, float)):
        return f"{value}".strip()
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        for item in value:
            text = _parse_water_level_text(item)
            if text:
                return text
        return ""
    if isinstance(value, dict):
        for key in ("text", "name", "value"):
            text = _parse_water_level_text(value.get(key))
            if text:
                return text
    return str(value).strip()


def _parse_bitable_from_url(url: str) -> tuple[Optional[str], Optional[str]]:
    if not url:
        return None, None
    parsed = urlparse(url)
    app_token = parsed.path.strip("/").split("/")[-1] if parsed.path else None
    query = parse_qs(parsed.query or "")
    table_id = (query.get("table") or [None])[0]
    return app_token, table_id


def fetch_water_level_from_feishu(target_date: date) -> tuple[Dict[str, Any], Optional[str]]:
    result = {"water_level": "--", "source_date": None}
    app_id = os.getenv("FEISHU_APP_ID", "").strip()
    app_secret = os.getenv("FEISHU_APP_SECRET", "").strip()
    app_token = os.getenv("FEISHU_BITABLE_APP_TOKEN", "").strip()
    table_id = os.getenv("FEISHU_BITABLE_TABLE_ID", "").strip()
    date_field = os.getenv("FEISHU_BITABLE_DATE_FIELD", "观测日期").strip() or "观测日期"
    water_field = os.getenv("FEISHU_BITABLE_WATER_LEVEL_FIELD", "水位高程").strip() or "水位高程"
    if (not app_token or not table_id) and os.getenv("FEISHU_BITABLE_URL"):
        p_app_token, p_table_id = _parse_bitable_from_url(os.getenv("FEISHU_BITABLE_URL", ""))
        app_token = app_token or (p_app_token or "")
        table_id = table_id or (p_table_id or "")

    if not app_id or not app_secret or not app_token or not table_id:
        return result, "飞书参数未完整配置，已使用默认水位"

    try:
        token_resp = requests.post(
            "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
            json={"app_id": app_id, "app_secret": app_secret},
            timeout=20,
        )
        token_resp.raise_for_status()
        token_data = token_resp.json()
        if token_data.get("code") != 0:
            raise ValueError(token_data.get("msg", "tenant_access_token 获取失败"))
        tenant_access_token = token_data.get("tenant_access_token")
        if not tenant_access_token:
            raise ValueError("tenant_access_token 为空")

        headers = {"Authorization": f"Bearer {tenant_access_token}"}
        page_token = None
        rows = []
        while True:
            params = {"page_size": 500}
            if page_token:
                params["page_token"] = page_token
            resp = requests.get(
                f"https://open.feishu.cn/open-apis/bitable/v1/apps/{app_token}/tables/{table_id}/records",
                headers=headers,
                params=params,
                timeout=25,
            )
            resp.raise_for_status()
            data = resp.json()
            if data.get("code") != 0:
                raise ValueError(data.get("msg", "bitable records 获取失败"))
            payload = data.get("data") or {}
            for item in payload.get("items", []):
                fields = item.get("fields") or {}
                obs_date = _parse_bitable_date(fields.get(date_field))
                water_text = _parse_water_level_text(fields.get(water_field))
                if water_text:
                    rows.append({"date": obs_date, "water_level": water_text})
            if not payload.get("has_more"):
                break
            page_token = payload.get("page_token")
            if not page_token:
                break

        if not rows:
            return result, "飞书未查询到水位记录，已使用默认水位"

        dated_rows = [x for x in rows if isinstance(x.get("date"), date)]
        selected = None
        if dated_rows:
            valid = [x for x in dated_rows if x["date"] <= target_date]
            selected = max(valid, key=lambda x: x["date"]) if valid else max(dated_rows, key=lambda x: x["date"])
        if selected is None:
            selected = rows[0]

        result["water_level"] = selected.get("water_level") or "--"
        result["source_date"] = selected.get("date")
        return result, None
    except Exception as e:
        return result, f"飞书水位获取失败，已使用默认水位: {e}"


def update_first_table_summary(
    doc: Document,
    date_text: str,
    weather_text: str,
    wind_text: str,
    water_level_text: str,
    temp_text: str,
) -> bool:
    if not doc.tables:
        return False
    table = doc.tables[0]
    if len(table.rows) < 2 or len(table.rows[1].cells) < 5:
        return False
    values = [date_text, weather_text, wind_text, water_level_text, temp_text]
    row = table.rows[1]
    for idx, value in enumerate(values):
        set_cell_text_preserve_style(row.cells[idx], value)
    return True


def set_cell_text_preserve_style(cell, text: str):
    """
    不破坏模板单元格样式地写入文本：
    - 保留首段落（继承对齐、间距、行距）
    - 仅替换首 run 文本（继承字体、字号等）
    """
    text = "" if text is None else str(text)

    while len(cell.paragraphs) > 1:
        extra = cell.paragraphs[-1]
        extra._element.getparent().remove(extra._element)

    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    if p.runs:
        p.runs[0].text = text
        for run in list(p.runs[1:]):
            p._element.remove(run._element)
    else:
        p.add_run(text)


def set_paragraph_text_preserve_style(paragraph, text: str):
    value = "" if text is None else str(text)
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in list(paragraph.runs[1:]):
            paragraph._element.remove(run._element)
    else:
        paragraph.add_run(value)


def update_paragraph_date(paragraph, target_day: date, allow_month_dot_day: bool = False) -> bool:
    if paragraph.runs:
        raw_text = "".join(run.text for run in paragraph.runs)
    else:
        raw_text = paragraph.text or ""
    new_text, changed = replace_date_in_text(raw_text, target_day, allow_month_dot_day=allow_month_dot_day)
    if changed:
        set_paragraph_text_preserve_style(paragraph, new_text)
    return changed


def update_table_date_by_index(doc: Document, table_index: int, target_day: date) -> bool:
    if table_index < 0 or table_index >= len(doc.tables):
        return False
    changed = False
    table = doc.tables[table_index]
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if update_paragraph_date(paragraph, target_day):
                    changed = True
    return changed


def update_table3_date(doc: Document, target_day: date, table_index: int = 2) -> bool:
    if table_index < 0 or table_index >= len(doc.tables):
        return False
    table = doc.tables[table_index]
    if len(table.rows) < 2 or len(table.rows[1].cells) < 1:
        return False

    changed = False
    date_cell = table.rows[1].cells[0]
    for paragraph in date_cell.paragraphs:
        if update_paragraph_date(paragraph, target_day, allow_month_dot_day=True):
            changed = True
    return changed


def update_body_signature_dates(doc: Document, target_day: date) -> bool:
    changed = False
    for paragraph in doc.paragraphs:
        if paragraph.runs:
            raw_text = "".join(run.text for run in paragraph.runs)
        else:
            raw_text = paragraph.text or ""
        if not _SIGNATURE_DATE_LABEL_REGEX.search(raw_text):
            continue
        new_text, replaced = replace_date_in_text(raw_text, target_day, allow_month_dot_day=True)
        if replaced:
            set_paragraph_text_preserve_style(paragraph, new_text)
            changed = True
    return changed


def update_footer_dates(doc: Document, target_day: date) -> bool:
    changed = False
    seen_footer_nodes: set[int] = set()
    for section in doc.sections:
        for attr in ("footer", "first_page_footer", "even_page_footer"):
            footer = getattr(section, attr, None)
            if footer is None:
                continue
            node_id = id(footer._element)
            if node_id in seen_footer_nodes:
                continue
            seen_footer_nodes.add(node_id)

            for paragraph in footer.paragraphs:
                if update_paragraph_date(paragraph, target_day):
                    changed = True

            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if update_paragraph_date(paragraph, target_day):
                                changed = True
    return changed


def clear_vmerge(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for vm in tc_pr.findall(qn('w:vMerge')):
        tc_pr.remove(vm)


def set_vmerge(cell, state: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    for vm in tc_pr.findall(qn('w:vMerge')):
        tc_pr.remove(vm)
    vm = OxmlElement('w:vMerge')
    vm.set(qn('w:val'), state)
    tc_pr.append(vm)


def apply_vertical_merge(table, data_start_row: int, items: List[Dict[str, str]], col_index: int, key_fn):
    # 先清理旧合并标记
    for i in range(len(items)):
        clear_vmerge(table.rows[data_start_row + i].cells[col_index])

    group_start = 0
    while group_start < len(items):
        current_key = key_fn(items[group_start])
        group_end = group_start + 1
        while group_end < len(items) and key_fn(items[group_end]) == current_key:
            group_end += 1

        if group_end - group_start > 1:
            top_cell = table.rows[data_start_row + group_start].cells[col_index]
            set_vmerge(top_cell, "restart")
            for idx in range(group_start + 1, group_end):
                c = table.rows[data_start_row + idx].cells[col_index]
                set_cell_text_preserve_style(c, "")
                set_vmerge(c, "continue")

        group_start = group_end


def render_daily_stats_table(doc: Document, items: List[Dict[str, str]]):
    table, header_row_index, header_map = find_daily_stats_table(doc)
    if table is None:
        raise ValueError(
            "未找到“当日施工统计表”（表头需包含："
            "序号/施工部位/施工内容/日完成量/备注 或 "
            "S/N/Construction Area/Activities/Quantities Completed/Remarks）"
        )

    data_start_row = header_row_index + 1
    style_row_index = data_start_row if len(table.rows) > data_start_row else header_row_index
    style_row_xml = deepcopy(table.rows[style_row_index]._tr)

    # 删除旧数据行（含多余空行），只保留表头
    for row_idx in range(len(table.rows) - 1, header_row_index, -1):
        table._tbl.remove(table.rows[row_idx]._tr)

    # 根据入参动态新增数据行
    for _ in items:
        table._tbl.append(deepcopy(style_row_xml))

    seq_col = header_map["seq"]
    location_col = header_map["location"]
    content_col = header_map["content"]
    quantity_col = header_map["quantity"]
    remarks_col = header_map["remarks"]

    # 写入数据（样式继承自模板行）
    for idx, item in enumerate(items):
        row = table.rows[data_start_row + idx]
        set_cell_text_preserve_style(row.cells[seq_col], item["seq"])
        set_cell_text_preserve_style(row.cells[location_col], item["location"])
        set_cell_text_preserve_style(row.cells[content_col], item["content"])
        set_cell_text_preserve_style(row.cells[quantity_col], item["quantity"])
        set_cell_text_preserve_style(row.cells[remarks_col], item["remarks"])

    # 纵向合并：支持“一个序号对应多个施工内容”
    apply_vertical_merge(
        table,
        data_start_row,
        items,
        seq_col,
        key_fn=lambda x: x["seq"],
    )

    # 同步合并施工部位（与序号分组一致时保持模板阅读习惯）
    apply_vertical_merge(
        table,
        data_start_row,
        items,
        location_col,
        key_fn=lambda x: (x["seq"], x["location"]),
    )

# ---------------- 模型定义 ----------------

class FillTemplateRequest(BaseModel):
    template_base64: str
    content: str
    daily_stats: Optional[List[Dict[str, Any]]] = None
    daily_stats_base64: Optional[str] = None
    strict_unit_check: bool = False
    table_index: int = 0
    row_index: int = 4
    col_index: int = 2
    update_date_weather: bool = False
    upload_to_feishu: bool = False
    feishu_token: Optional[str] = None


class GenerateFromTemplateRequest(BaseModel):
    daily_stats_base64: Optional[str] = None
    chinese_data: Optional[str] = None
    english_data: Optional[str] = None
    cn_template_base64: Optional[str] = None
    en_template_base64: Optional[str] = None
    trigger_date: Optional[str] = None

class UpdateDateWeatherRequest(BaseModel):
    document_base64: str
    feishu_token: Optional[str] = None

class UpdatePersonnelRequest(BaseModel):
    document_base64: str
    personnel_text: str 
    feishu_token: Optional[str] = None

class AppendixTableData(BaseModel):
    table_index: int
    row_name: str
    today_qty: str
    total_qty: str

class UpdateAppendixRequest(BaseModel):
    document_base64: str
    data: List[AppendixTableData]
    feishu_token: Optional[str] = None

# ---------------- API 接口实现 ----------------

@app.get("/")
async def root():
    return {"success": True, "message": "Word service is running"}


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/generate-from-template")
async def generate_from_template(req: GenerateFromTemplateRequest):
    try:
        def load_template_doc(template_b64: Optional[str], fallback_path: str, label: str) -> Document:
            if template_b64:
                try:
                    raw = base64.b64decode(template_b64)
                    return Document(io.BytesIO(raw))
                except Exception as exc:
                    raise ValueError(f"{label}_template_base64 解析失败，请传入有效的 docx Base64") from exc
            if os.path.exists(fallback_path):
                return Document(fallback_path)
            raise FileNotFoundError(f"未提供 {label}_template_base64，且找不到本地模板文件：{fallback_path}")

        warnings: List[str] = []
        trigger_day = parse_trigger_date(req.trigger_date)

        cn_items = parse_daily_stats_from_base64(req.daily_stats_base64 or "")
        if cn_items is None:
            cn_items = parse_daily_stats_json_text(req.chinese_data)
        if cn_items is None:
            cn_items = []
        if not cn_items:
            raise ValueError("未解析到可用的中文施工统计数据")

        bad_rows = detect_suspect_quantity(cn_items)
        if bad_rows:
            return {"success": False, "message": f"检测到疑似损坏单位（如 m?），问题行: {bad_rows}"}
        if detect_garbled_daily_stats(cn_items):
            raise ValueError("检测到疑似乱码，请确保 daily_stats_base64/chinese_data 使用 UTF-8")

        en_items = parse_english_translated_items(req.english_data)
        if en_items is None:
            en_items = cn_items

        weather_data, weather_warning = fetch_pakbeng_weather(trigger_day)
        if weather_warning:
            warnings.append(weather_warning)
        water_data, water_warning = fetch_water_level_from_feishu(trigger_day)
        if water_warning:
            warnings.append(water_warning)

        cn_date_text = format_cn_date(trigger_day)
        en_date_text = format_en_date(trigger_day)
        water_level_text = str(water_data.get("water_level") or "--")
        water_source_date = water_data.get("source_date")

        cn_doc = load_template_doc(req.cn_template_base64, CN_TEMPLATE_PATH, "cn")
        render_daily_stats_table(cn_doc, cn_items)
        update_first_table_summary(
            cn_doc,
            date_text=cn_date_text,
            weather_text=weather_data.get("weather_zh", "未知"),
            wind_text=weather_data.get("wind_zh", "未知"),
            water_level_text=water_level_text,
            temp_text=weather_data.get("temp", "--"),
        )
        update_table3_date(cn_doc, trigger_day)
        update_body_signature_dates(cn_doc, trigger_day)
        update_footer_dates(cn_doc, trigger_day)
        cn_out = io.BytesIO()
        cn_doc.save(cn_out)

        en_doc = load_template_doc(req.en_template_base64, EN_TEMPLATE_PATH, "en")
        render_daily_stats_table(en_doc, en_items)
        update_first_table_summary(
            en_doc,
            date_text=en_date_text,
            weather_text=weather_data.get("weather_en", "Unknown"),
            wind_text=weather_data.get("wind_en", "Unknown"),
            water_level_text=water_level_text,
            temp_text=weather_data.get("temp", "--"),
        )
        update_table3_date(en_doc, trigger_day)
        update_body_signature_dates(en_doc, trigger_day)
        update_footer_dates(en_doc, trigger_day)
        en_out = io.BytesIO()
        en_doc.save(en_out)

        result = {
            "success": True,
            "cn_document_base64": base64.b64encode(cn_out.getvalue()).decode(),
            "en_document_base64": base64.b64encode(en_out.getvalue()).decode(),
            "weather_info": {
                "date": cn_date_text,
                "weather": weather_data.get("weather_zh", "未知"),
                "temp": weather_data.get("temp", "--"),
                "wind": weather_data.get("wind_zh", "未知"),
                "water_level": water_level_text,
                "water_level_date": format_cn_date(water_source_date) if isinstance(water_source_date, date) else "",
            },
        }
        if warnings:
            result["warnings"] = warnings
        return result
    except Exception as e:
        return {"success": False, "message": str(e)}


@app.post("/fill-template")
async def fill_template(req: FillTemplateRequest):
    try:
        file_bytes = base64.b64decode(req.template_base64)
        doc = Document(io.BytesIO(file_bytes))
        warnings: List[str] = []

        # 模式1：当日施工统计表动态更新（优先）
        use_daily_stats_mode = False
        normalized_items: List[Dict[str, str]] = []
        if req.daily_stats_base64:
            use_daily_stats_mode = True
            parsed_b64 = parse_daily_stats_from_base64(req.daily_stats_base64)
            if parsed_b64 is None:
                raise ValueError("daily_stats_base64 解析失败，请传 UTF-8 JSON 数组的 Base64")
            normalized_items = parsed_b64
        elif req.daily_stats is not None:
            use_daily_stats_mode = True
            normalized_items = normalize_daily_stats_items(req.daily_stats)
        else:
            parsed = parse_daily_stats_from_content(req.content)
            if parsed is not None:
                use_daily_stats_mode = True
                normalized_items = parsed

        if use_daily_stats_mode:
            bad_quantity_rows = detect_suspect_quantity(normalized_items)
            if bad_quantity_rows:
                if req.strict_unit_check:
                    raise ValueError(
                        "检测到损坏单位（如 m? / m�），已拒绝渲染。"
                        f"请修正后重试，问题行: {bad_quantity_rows}"
                    )
                warnings.append(
                    "检测到损坏单位（如 m? / m�），已按原值保留并继续渲染。"
                    f"问题行: {bad_quantity_rows}"
                )
            if detect_garbled_daily_stats(normalized_items):
                raise ValueError(
                    "检测到疑似乱码（内容出现 ??? 且无中文）。请使用 UTF-8 JSON，"
                    "或改用 daily_stats_base64 传入（推荐）。"
                )
            render_daily_stats_table(doc, normalized_items)
        else:
            # 模式2：兼容旧版单单元格逐行排版
            if doc.tables and len(doc.tables) > req.table_index:
                table = doc.tables[req.table_index]
                if len(table.rows) > req.row_index:
                    cell = table.cell(req.row_index, req.col_index)

                    # 清空单元格
                    cell.text = ""

                    # 逐行处理，精确控制格式
                    lines = req.content.split('\n')
                    for line in lines:
                        process_and_add_line(cell, line)
        
        out = io.BytesIO()
        doc.save(out)
        result = {"success": True, "document_base64": base64.b64encode(out.getvalue()).decode()}
        if warnings:
            result["warnings"] = warnings
        return result
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "message": str(e)}

@app.post("/update-date-weather")
async def update_date_weather(req: UpdateDateWeatherRequest):
    try:
        file_bytes = base64.b64decode(req.document_base64)
        doc = Document(io.BytesIO(file_bytes))
        
        trigger_day = parse_trigger_date(None)
        date_str = format_cn_date(trigger_day)
        weather_data, _ = fetch_pakbeng_weather(trigger_day)
        weather_str = (
            f"天气：{weather_data.get('weather_zh', '未知')}"
            f"                气温：{weather_data.get('temp', '--')}"
        )
        
        if doc.tables:
            table = doc.tables[0]
            if len(table.rows) > 0:
                cells = table.rows[0].cells
                if len(cells) > 0: 
                    cells[0].text = ""
                    run = cells[0].paragraphs[0].add_run(date_str)
                    format_run_font(run)
                if len(cells) > 1: 
                    cells[-1].text = ""
                    run = cells[-1].paragraphs[0].add_run(weather_str)
                    format_run_font(run)

        update_table3_date(doc, trigger_day)
        update_body_signature_dates(doc, trigger_day)
        update_footer_dates(doc, trigger_day)
        
        out = io.BytesIO()
        doc.save(out)
        return {"success": True, "document_base64": base64.b64encode(out.getvalue()).decode()}
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.post("/update-personnel-stats")
async def update_personnel_stats(req: UpdatePersonnelRequest):
    try:
        file_bytes = base64.b64decode(req.document_base64)
        doc = Document(io.BytesIO(file_bytes))
        
        # 统计信息在文末追加，默认不需要特殊缩进，但需要字体规范
        p = doc.add_paragraph()
        run = p.add_run("\n" + req.personnel_text)
        format_run_font(run)
        
        out = io.BytesIO()
        doc.save(out)
        return {"success": True, "document_base64": base64.b64encode(out.getvalue()).decode()}
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.post("/update-appendix-tables")
async def update_appendix_tables(req: UpdateAppendixRequest):
    try:
        file_bytes = base64.b64decode(req.document_base64)
        doc = Document(io.BytesIO(file_bytes))
        
        for item in req.data:
            target_table = find_target_table(doc, item.table_index)
            if target_table:
                update_table_row(target_table, item.row_name, item.today_qty, item.total_qty)
        
        out = io.BytesIO()
        doc.save(out)
        return {"success": True, "document_base64": base64.b64encode(out.getvalue()).decode()}
    except Exception as e:
        return {"success": False, "message": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
